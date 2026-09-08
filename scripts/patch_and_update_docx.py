"""Patch and update SecureMask research paper docx draft into final IEEE submission docx.

Applies:
  - Fixed [Content_Types].xml to eliminate corruption / undefined media KeyError
  - Reformed Two-Component PEI formulation and schema masking ratios mu_f
  - Section III Architecture updates (EasyOCR primary, PaddleOCR fallback, defensive QR, failure-safe redactor)
  - Section VII Experimental Results: Real N=3 human study correlation (r=0.9736, rho=0.8898, bootstrap CIs)
  - Synthetic benchmark evaluation results & Systematic Ablation Battery table
  - Elimination of all stale / contradictory text blocks (old w*2/w*10, PEI=20.0 bug as future work, N=5, 0.940/0.870, data collection in progress)
  - Embeds publication figures (fig1 to fig5) directly into the Word document

Usage::
    python scripts/patch_and_update_docx.py
"""
from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC_DOCX = Path(r"C:\Users\Atharv\Desktop\SecureMask_Research_Paper_Draft (5).docx")
OUT_DOCX = Path(r"C:\Users\Atharv\Desktop\SecureMask_Research_Paper_Draft_Final.docx")
FIG_DIR = Path("paper_figures")


def repair_and_load_docx(src_path: Path) -> Document:
    """Fix undefined content-type and load python-docx Document object."""
    fixed_buf = io.BytesIO()
    with zipfile.ZipFile(src_path, "r") as zin:
        with zipfile.ZipFile(fixed_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                content = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    content_str = content.decode("utf-8")
                    if 'Extension="undefined"' not in content_str:
                        content_str = content_str.replace(
                            '<Default ContentType="image/png" Extension="png"/>',
                            '<Default ContentType="image/png" Extension="png"/><Default ContentType="image/png" Extension="undefined"/>'
                        )
                        content = content_str.encode("utf-8")
                zout.writestr(item, content)

    fixed_buf.seek(0)
    return Document(fixed_buf)


def update_paper_content(doc: Document) -> None:
    """Update sections, text, formulas, and tables in the document."""

    # 1. Update Abstract (Paragraph 2)
    abstract_text = (
        "Abstract—Identity documents such as Aadhaar cards, PAN cards, passports, driving licences, "
        "and voter ID cards are routinely photographed and shared for verification purposes in India, "
        "frequently exposing sensitive personal data beyond what the receiving context requires, in "
        "violation of Section 6 of India's Digital Personal Data Protection Act (DPDP Act), 2023. "
        "Existing redaction tools apply rigid, purpose-agnostic heuristics that ignore the transaction "
        "purpose. This paper presents SecureMask, an end-to-end framework combining a dual-language "
        "(English/Hindi) OCR pipeline (EasyOCR primary with PaddleOCR fallback), a fine-tuned MobileNetV2 "
        "document classifier with calibrated keyword fallback, multi-modal field extraction (fuzzy regex, "
        "spaCy NER, and defensive QR/XML parsing), and a formally derived, two-component Privacy Exposure "
        "Index (PEI) on a 0–100 scale. PEI decomposes exposure into excess disclosure and contextually "
        "necessary primary identifier exposure, integrating schema-proportional masking factors and "
        "eliminating mathematical floor anomalies. On a pilot human-validation study (N = 3 real human raters, "
        "12 standardized scenarios), SecureMask's reformed PEI demonstrates exceptional alignment with human "
        "risk judgements (Pearson r = 0.9736, 95% CI [0.9507, 0.9955], p = 9.59e-8; Spearman rho = 0.8898, "
        "95% CI [0.5474, 0.9742], p = 1.06e-4). Controlled architectural ablations on a 50-document synthetic "
        "benchmark demonstrate that context-aware redaction reduces privacy exposure by 65.7 PEI points while "
        "preserving 100% required transactional utility, completely eliminating the 54.4% over-redaction and "
        "31.1% under-redaction induced by purpose-agnostic static baselines. Evaluation on the synthetic benchmark "
        "yields 86.0% classification accuracy and 0.6266 normalized extraction F1 across 50 annotated Indian identity credentials."
    )
    if len(doc.paragraphs) > 2:
        doc.paragraphs[2].text = abstract_text

    # 2. Update Section I overview paragraph (remove "empirical results are still being collected")
    for p in doc.paragraphs:
        if "since at the time of writing empirical results are still being collected" in p.text:
            p.text = (
                "We describe the architecture in Section III, the PEI formulation in Section IV, the necessity matrix "
                "in Section V, the explainability layer in Section VI, and comprehensive empirical evaluations—including an "
                "N = 3 real human perception study (E4), synthetic benchmark classification (E1) and extraction (E2), "
                "systematic architectural ablations, robustness analysis under perturbations (E6), and CPU latency "
                "profiling (E7)—in Section VII. Section VIII discusses limitations and ethical compliance, followed by "
                "concluding remarks in Section IX."
            )

    # 3. Update Section III Architecture details
    for p in doc.paragraphs:
        if "HuggingFace IndicNER" in p.text:
            p.text = (
                "Field values are located using a combination of regular expressions and fuzzy string matching (via RapidFuzz) "
                "anchored to nearby keyword tokens, so that OCR misreads of digits (e.g., a zero misread as the letter O) can "
                "still be recovered. Devanagari names are detected by merging consecutive Devanagari tokens, stripping a blacklist "
                "of institutional words, and applying length-based heuristics. Where regex and fuzzy matching are insufficient—particularly "
                "for free-form names and addresses—a named-entity recognition fallback using spaCy's en_core_web_sm is applied. "
                "Aadhaar QR codes are decoded (pyzbar) and their embedded, zlib-compressed XML cross-checked against the OCR-extracted "
                "fields, with DOCTYPE and ENTITY stripping to prevent XXE attacks and decompression caps (512 KB). Faces are located "
                "with Haar cascade classifiers and signatures via contour analysis and aspect-ratio heuristics, treated as high-risk biometric fields."
            )
        elif "covered with an opaque white rectangle" in p.text:
            p.text = (
                "Fields marked for full redaction are covered with an opaque black rectangle (solid zero-leakage overlay); fields "
                "marked for partial masking are covered with a black rectangle over their leading portion, leaving trailing characters "
                "(e.g., the last four digits of an identifier) visible for partial verification. To guarantee failure-safe rendering, "
                "bounding boxes are validated against non-positive coordinates and clamped to physical image boundaries [0, W] × [0, H], "
                "preventing out-of-bounds rendering exceptions. Redaction is applied at the pixel level with per-field configurable "
                "padding using Pillow, and the anonymised image is written to an isolated storage path, severing any link back to the unredacted original."
            )

    # 4. Update PEI Section (Formulas and Definitions)
    for p in doc.paragraphs:
        if "raw_score = " in p.text:
            p.text = (
                "PEI(D, c) = [ Σ_{f ∈ F_excess} (e_f · w_f) + λ · Σ_{f ∈ F_{id,req}} (e_f · w_f) ] "
                "/ [ Σ_{f ∈ F_all} w_f ] × 100"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "max_possible = " in p.text:
            p.text = (
                "where F_excess represents fields that are redundant (n_f = False) or unconditionally sensitive "
                "(always_redact = True); F_{id,req} denotes contextually required primary national identifiers "
                "(aadhaar_number, pan_number, passport_number, dl_number, epic_number); w_f is the field sensitivity weight; "
                "e_f ∈ {1.0 (allow), 0.0 (redact), μ_f (mask)} is the decision exposure factor; and λ = 0.50 is the "
                "calibrated policy parameter. The schema masking factor is determined by μ_f = (visible characters) / (total characters) "
                "(e.g., 4/12 for Aadhaar, 4/10 for PAN, 4/8 for Passport, 4/15 for Driving License, 4/10 for Voter ID). When only necessary atomic "
                "attributes (Name, DOB, Address) are disclosed, PEI evaluates to exactly 0.0, eliminating the mathematical floor."
            )

    # 5. Update Table IV (E4 Scenario Table) with Real Human Study Results
    if len(doc.tables) >= 4:
        table4 = doc.tables[3]
        # Real empirical data from E4_Human_Rating_Sheet.csv and compute_pei_details (lambda=0.50)
        e4_data = [
            ("Scenario 1", "Age Verification (Aadhaar - Unredacted)", "9.17", "71.9"),
            ("Scenario 2", "Minimal Age Verification (Aadhaar - Masked)", "2.10", "0.0"),
            ("Scenario 3", "Identity Verification (PAN - Unredacted)", "8.83", "54.8"),
            ("Scenario 4", "Identity Verification (PAN - Masked)", "3.33", "11.9"),
            ("Scenario 5", "Address Proof (DL - Unredacted)", "8.50", "61.4"),
            ("Scenario 6", "Address Proof (DL - Masked)", "2.77", "0.0"),
            ("Scenario 7", "KYC Onboarding (Passport - Unredacted)", "9.50", "54.2"),
            ("Scenario 8", "KYC Onboarding (Passport - Masked)", "4.17", "5.2"),
            ("Scenario 9", "General File Upload (Voter ID - Unredacted)", "9.83", "77.3"),
            ("Scenario 10", "General File Upload (Voter ID - Masked)", "1.83", "0.0"),
            ("Scenario 11", "Address Proof (Aadhaar - Unredacted)", "9.00", "73.4"),
            ("Scenario 12", "Address Proof (Aadhaar - Masked)", "2.00", "0.0"),
        ]

        for r_idx, row_data in enumerate(e4_data, start=1):
            if r_idx < len(table4.rows):
                cells = table4.rows[r_idx].cells
                if len(cells) >= 4:
                    cells[0].text = row_data[0]
                    cells[1].text = row_data[1]
                    cells[2].text = row_data[2]
                    cells[3].text = row_data[3]

    # 6. Update Section VII text (replace old N=5, 0.940, future work claims)
    for p in doc.paragraphs:
        if "Per-document-type classification, extraction, redaction, and latency results" in p.text:
            p.text = (
                "Document classification, field extraction, robustness, and latency benchmarks are evaluated on our annotated "
                "synthetic benchmark (50 documents, 10 per credential category; storage/synthetic_benchmark). E1 document classification "
                "achieves 86.0% overall accuracy across the five classes (Aadhaar 100%, PAN 100%, DL 100%, Passport F1 0.7407, Voter ID F1 0.4615). "
                "E2 field extraction achieves a normalized F1 of 0.6266 (strict F1 0.6116). Full architectural ablations confirm that RapidFuzz "
                "fuzzy token alignment yields a +23.07% absolute recall improvement over exact matching (0.8769 vs 0.6462)."
            )
        elif "E4 specifically has since been carried out. Five independent raters" in p.text or "Five independent raters, unfamiliar" in p.text:
            p.text = (
                "For E4 specifically, an empirical pilot human-validation study was conducted with N = 3 real human raters across 12 "
                "standardized paired scenarios (six document/context pairs evaluated under naive unredacted disclosure versus SecureMask "
                "context-aware redaction). Raters scored perceived privacy risk on a 1–10 scale. Benchmarking these scores against SecureMask's "
                "reformed PEI (λ = 0.50) yields exceptional agreement: Pearson linear correlation r = 0.9736 (p = 9.59e-8, 95% bootstrap CI "
                "[0.9507, 0.9955], R² = 0.9479) and Spearman rank correlation rho = 0.8898 (p = 1.06e-4, 95% bootstrap CI [0.5474, 0.9742]). "
                "Sensitivity analysis demonstrates that inter-tier rank ordering across unredacted, partially masked, and minimal tiers remains "
                "strictly preserved across λ ∈ [0.25, 1.00], with discriminability margin peaking at λ = 0.50 (62.6 points)."
            )

    # 7. Update Section VIII Discussion (replace old OCR list and floor effect future work)
    for p in doc.paragraphs:
        if "PaddleOCR, EasyOCR, IndicNER, spaCy" in p.text:
            p.text = (
                "Third, OCR and NER components—specifically EasyOCR (primary dual-language text detector and recognizer), PaddleOCR "
                "(fallback OCR engine), and spaCy NER—are pretrained models; their error characteristics on Indian document layouts "
                "are evaluated on our 50-document synthetic benchmark in Section VII."
            )
        elif "Fifth, the E4 human-validation study" in p.text and "floor effect" in p.text:
            p.text = (
                "Fifth, during the early development audit of the initial prototype, an algebraic floor effect was identified in the "
                "historical single-penalty calculation: computing the denominator over only surviving allowed fields caused an invariant "
                "floor of PEI = 20.0 whenever all surviving fields were necessary, while treating partial masking identically to full "
                "redaction ignored residual identifier risk. In the reformed SecureMask architecture presented here, this limitation has "
                "been formally resolved: the denominator is normalized across all detected fields in the document, and partial masking is "
                "weighted by the schema-derived character ratio μ_f. As validated in Section VII and Table IV, the 20.0 floor is eliminated "
                "(evaluating to exactly 0.0 for minimal disclosures), and partial masking of primary identifiers properly registers bounded "
                "residual exposure (e.g., PEI = 11.9 for masked PAN, 5.2 for masked Passport), yielding the verified correlation of r = 0.9736 "
                "with human risk judgements."
            )

    # 8. Update Section IX Conclusion
    for p in doc.paragraphs:
        if "the immediate next step is executing the evaluation protocol" in p.text:
            p.text = (
                "This paper presented SecureMask, an end-to-end context-aware privacy preservation framework for Indian identity "
                "credentials. By replacing rigid, purpose-agnostic blackout heuristics with a formally grounded, two-component Privacy "
                "Exposure Index (PEI) and a deterministic necessity matrix, SecureMask enables automated data minimization aligned with "
                "India's DPDPA 2023. Empirical validation against real human privacy raters (r = 0.9736) and systematic architectural "
                "ablations on a 50-document synthetic benchmark confirm that SecureMask eliminates excessive disclosure while guaranteeing "
                "that essential transaction credentials remain verifiable."
            )

    # 9. Embed / Update Publication Figures
    print("Embedding publication figures into document...")

    # Fig 1: Architecture
    fig1_path = FIG_DIR / "fig1_architecture.png"
    if fig1_path.exists():
        for p in doc.paragraphs:
            if "Fig. 1." in p.text and len(p.runs) > 0:
                p_img = p.insert_paragraph_before()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(str(fig1_path), width=Inches(6.5))
                break

    # Fig 2: E4 Correlation
    fig2_path = FIG_DIR / "fig2_e4_correlation.png"
    if fig2_path.exists():
        for p in doc.paragraphs:
            if "TABLE IV." in p.text:
                p_img = p.insert_paragraph_before()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(str(fig2_path), width=Inches(5.5))
                p_caption = p.insert_paragraph_before("Fig. 2. SecureMask PEI vs. Human Risk Perception (N=3 Real Raters, 95% Bootstrap CI).")
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # Fig 3: Lambda Sensitivity
    fig3_path = FIG_DIR / "fig3_lambda_sensitivity.png"
    if fig3_path.exists():
        for p in doc.paragraphs:
            if "VIII. DISCUSSION" in p.text:
                p_img = p.insert_paragraph_before()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(str(fig3_path), width=Inches(6.0))
                p_caption = p.insert_paragraph_before("Fig. 3. Sensitivity Analysis of Policy Parameter λ: Correlation & Tier Discriminability.")
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # Fig 4: Robustness
    fig4_path = FIG_DIR / "fig4_robustness.png"
    if fig4_path.exists():
        for p in doc.paragraphs:
            if "E6. Robustness" in p.text or "VIII. DISCUSSION" in p.text:
                p_img = p.insert_paragraph_before()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(str(fig4_path), width=Inches(5.8))
                p_caption = p.insert_paragraph_before("Fig. 4. Robustness Degradation Curves Across Visual Perturbations (Blur, Skew, Illumination, JPEG).")
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # Fig 5: Latency
    fig5_path = FIG_DIR / "fig5_latency.png"
    if fig5_path.exists():
        for p in doc.paragraphs:
            if "IX. CONCLUSION" in p.text:
                p_img = p.insert_paragraph_before()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(str(fig5_path), width=Inches(5.5))
                p_caption = p.insert_paragraph_before("Fig. 5. Component-Wise Latency Profile on CPU (Total Mean: 7678.1 ms, P95: 11185.7 ms, Single Thread).")
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # 10. Update or add Ablation Summary Table callout in Section VII
    for p in doc.paragraphs:
        if "TABLE IV." in p.text:
            callout = p.insert_paragraph_before(
                "TABLE V. SYSTEMATIC ABLATION BATTERY: PRIVACY & UTILITY TRADEOFFS (SYNTHETIC BENCHMARK)\n"
                "• Full SecureMask: Pre-PEI = 67.3, Post-PEI = 1.7, ΔPEI = 65.7, Utility = 100.0%, Over-Redaction = 0.0%, Under-Redaction = 0.0%\n"
                "• Ablation A (Static Baseline): Pre-PEI = 67.3, Post-PEI = 20.3, ΔPEI = 47.1, Utility = 45.6%, Over-Redaction = 54.4%, Under-Redaction = 31.1%\n"
                "• Ablation B (Unweighted Ratio): Pre-PEI = 100.0, Post-PEI = 37.7, ΔPEI = 62.3, Utility = 100.0%, Over-Redaction = 0.0%, Under-Redaction = 0.0%\n"
                "• Ablation E (Binary Redaction): Pre-PEI = 67.3, Post-PEI = 3.7, ΔPEI = 63.6, Utility = 100.0%, Over-Redaction = 0.0%, Under-Redaction = 0.0%\n"
                "• Extraction Ablations: Full F1 = 0.9344 (Recall = 0.8769); Without Fuzzy Matching F1 = 0.7850 (Recall = 0.6462, -23.07% recall penalty)."
            )
            callout.alignment = WD_ALIGN_PARAGRAPH.LEFT
            break


def main():
    print(f"Loading source docx: {SRC_DOCX}")
    doc = repair_and_load_docx(SRC_DOCX)
    print("Updating content, formulas, empirical tables, and figures...")
    update_paper_content(doc)
    print(f"Saving final IEEE publication docx to: {OUT_DOCX}")
    doc.save(OUT_DOCX)
    print("Successfully generated final publication docx!")


if __name__ == "__main__":
    main()
