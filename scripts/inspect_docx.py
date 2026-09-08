import zipfile
import io
from pathlib import Path
from docx import Document

docx_path = Path(r"C:\Users\Atharv\Desktop\SecureMask_Research_Paper_Draft (5).docx")

# Read zip and inject .undefined into [Content_Types].xml
fixed_buf = io.BytesIO()
with zipfile.ZipFile(docx_path, 'r') as zin:
    with zipfile.ZipFile(fixed_buf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == '[Content_Types].xml':
                content_str = content.decode('utf-8')
                if 'Extension="undefined"' not in content_str:
                    content_str = content_str.replace(
                        '<Default ContentType="image/png" Extension="png"/>',
                        '<Default ContentType="image/png" Extension="png"/><Default ContentType="image/png" Extension="undefined"/>'
                    )
                    content = content_str.encode('utf-8')
            zout.writestr(item, content)

fixed_buf.seek(0)
doc = Document(fixed_buf)
print(f"Successfully loaded docx! Total paragraphs: {len(doc.paragraphs)}, Total tables: {len(doc.tables)}")

# Print first 15 paragraphs
for i, p in enumerate(doc.paragraphs[:15]):
    if p.text.strip():
        print(f"P[{i}]: {p.text[:90]}...")
